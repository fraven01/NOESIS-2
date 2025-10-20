from __future__ import annotations

import base64
import hashlib
import io
import zipfile

import pytest
from docx import Document
from docx.oxml import OxmlElement

from documents.contracts import InlineBlob
from documents.parsers_docx import DocxDocumentParser


_PNG_BASE64 = "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/w8AAwMCAO+X8a0AAAAASUVORK5CYII="

_NS = {
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
}


class _DocumentStub:
    def __init__(self, media_type: str, blob: InlineBlob) -> None:
        self.media_type = media_type
        self.blob = blob


def _inline_blob_from_payload(payload: bytes) -> InlineBlob:
    encoded = base64.b64encode(payload).decode("ascii")
    digest = hashlib.sha256(payload).hexdigest()
    return InlineBlob(
        type="inline",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        base64=encoded,
        sha256=digest,
        size=len(payload),
    )


def _set_inline_alt_text(run, description: str) -> None:
    """Inject descriptive metadata for the given run's inline image."""

    for inline in run._element.iter(f"{{{_NS['wp']}}}inline"):
        doc_pr = inline.find(f"{{{_NS['wp']}}}docPr")
        if doc_pr is None:
            doc_pr = OxmlElement("wp:docPr")
            doc_pr.set("id", "1")
            doc_pr.set("name", "Picture")
            inline.insert(0, doc_pr)
        doc_pr.set("descr", description)
        if not doc_pr.get("name"):
            doc_pr.set("name", "Picture")


def _save_document(document: Document) -> bytes:
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _build_docx_payload() -> tuple[bytes, bytes]:
    image_bytes = base64.b64decode(_PNG_BASE64)
    document = Document()
    document.core_properties.language = "en-US"
    document.add_heading("Project Plan", level=1)
    document.add_paragraph("Introduction paragraph.")
    document.add_paragraph("First bullet item", style="List Bullet")

    table = document.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "Header A"
    table.rows[0].cells[1].text = "Header B"
    table.rows[1].cells[0].text = "Row1 Col1"
    table.rows[1].cells[1].text = "Row1 Col2"
    table.rows[2].cells[0].text = "Row2 Col1"
    table.rows[2].cells[1].text = "Row2 Col2"

    paragraph = document.add_paragraph()
    paragraph.add_run("Before image text.")
    run = paragraph.add_run()
    run.add_picture(io.BytesIO(image_bytes))
    _set_inline_alt_text(run, "An illustrative screenshot")
    paragraph.add_run("After image text.")

    document.add_paragraph("Closing paragraph.")

    return _save_document(document), image_bytes


def _build_docx_payload_with_alt_fallback() -> tuple[bytes, bytes]:
    image_bytes = base64.b64decode(_PNG_BASE64)
    document = Document()
    document.core_properties.language = "en-US"
    paragraph = document.add_paragraph("Image intro text.")
    run = paragraph.add_run()
    run.add_picture(io.BytesIO(image_bytes))
    _set_inline_alt_text(run, "Fallback description")
    return _save_document(document), image_bytes


def _build_docx_payload_missing_image() -> bytes:
    payload, _ = _build_docx_payload()
    source = io.BytesIO(payload)
    buffer = io.BytesIO()
    with zipfile.ZipFile(source) as src, zipfile.ZipFile(buffer, "w") as dst:
        for info in src.infolist():
            if info.filename.startswith("word/media/"):
                continue
            dst.writestr(info, src.read(info.filename))
    return buffer.getvalue()


def _build_docx_payload_with_large_image(byte_count: int) -> bytes:
    payload, _ = _build_docx_payload()
    source = io.BytesIO(payload)
    buffer = io.BytesIO()
    with zipfile.ZipFile(source) as src, zipfile.ZipFile(buffer, "w") as dst:
        for info in src.infolist():
            data = src.read(info.filename)
            if info.filename.startswith("word/media/"):
                data = b"0" * byte_count
            dst.writestr(info, data)
    return buffer.getvalue()


def _build_docx_payload_exceeding_entry_limit(extra_entries: int) -> bytes:
    payload, _ = _build_docx_payload()
    source = io.BytesIO(payload)
    buffer = io.BytesIO()
    with zipfile.ZipFile(source) as src, zipfile.ZipFile(buffer, "w") as dst:
        for info in src.infolist():
            dst.writestr(info, src.read(info.filename))
        for index in range(extra_entries):
            dst.writestr(f"word/media/filler{index}.bin", b"")
    return buffer.getvalue()


def _build_docx_payload_with_generic_image_type() -> bytes:
    payload, _ = _build_docx_payload()
    source = io.BytesIO(payload)
    buffer = io.BytesIO()
    with zipfile.ZipFile(source) as src, zipfile.ZipFile(buffer, "w") as dst:
        for info in src.infolist():
            data = src.read(info.filename)
            if info.filename == "[Content_Types].xml":
                text = data.decode("utf-8")
                text = text.replace(
                    'ContentType="image/png"',
                    'ContentType="application/octet-stream"',
                )
                dst.writestr(info, text)
            else:
                dst.writestr(info, data)
    return buffer.getvalue()


def test_docx_parser_extracts_text_blocks_and_assets() -> None:
    payload, image_bytes = _build_docx_payload()
    blob = _inline_blob_from_payload(payload)
    document = _DocumentStub(
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        blob=blob,
    )
    parser = DocxDocumentParser()

    assert parser.can_handle(document)

    result = parser.parse(document, config={})

    heading_blocks = [block for block in result.text_blocks if block.kind == "heading"]
    paragraph_blocks = [
        block for block in result.text_blocks if block.kind == "paragraph"
    ]
    list_blocks = [block for block in result.text_blocks if block.kind == "list"]
    table_blocks = [
        block for block in result.text_blocks if block.kind == "table_summary"
    ]

    assert heading_blocks and heading_blocks[0].text == "Project Plan"
    assert heading_blocks[0].section_path == ("Project Plan",)

    assert any(block.text == "Introduction paragraph." for block in paragraph_blocks)
    assert list_blocks and list_blocks[0].text == "First bullet item"

    assert table_blocks and "rows=2" in table_blocks[0].text
    assert table_blocks[0].table_meta == {
        "rows": 2,
        "columns": 2,
        "headers": ["Header A", "Header B"],
        "sample_row_count": 2,
        "sample_rows": [["Row1 Col1", "Row1 Col2"], ["Row2 Col1", "Row2 Col2"]],
    }

    assert result.assets and len(result.assets) == 1
    asset = result.assets[0]
    assert asset.media_type == "image/png"
    assert asset.content == image_bytes
    assert asset.context_before and "Before image text." in asset.context_before
    assert asset.context_after and "After image text." in asset.context_after
    assert len(asset.context_before) <= 512
    assert len(asset.context_after) <= 512
    assert asset.metadata.get("parent_ref", "").startswith("paragraph:")
    assert asset.metadata.get("locator", "").startswith(asset.metadata["parent_ref"])
    candidates = asset.metadata.get("caption_candidates", [])
    assert candidates and candidates[0] == ("alt_text", "An illustrative screenshot")
    assert {block.language for block in result.text_blocks} == {"en-US"}

    stats = result.statistics
    assert stats["parser.tables"] >= 1
    assert stats["assets.images"] == len(result.assets)
    assert stats["parser.words"] >= 1
    assert stats["parse.blocks.total"] == len(result.text_blocks)
    assert stats["parse.assets.total"] == len(result.assets)


def test_docx_parser_detects_by_blob_when_media_type_generic() -> None:
    payload, _ = _build_docx_payload()
    blob = _inline_blob_from_payload(payload)
    document = _DocumentStub(media_type="application/octet-stream", blob=blob)
    parser = DocxDocumentParser()

    assert parser.can_handle(document)


def test_docx_parser_falls_back_to_extension_for_unknown_media_type() -> None:
    payload = _build_docx_payload_with_generic_image_type()
    blob = _inline_blob_from_payload(payload)
    document = _DocumentStub(
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        blob=blob,
    )
    parser = DocxDocumentParser()

    result = parser.parse(document, config={})

    assert result.assets
    assert result.assets[0].media_type == "image/png"


def test_docx_parser_raises_for_missing_image_target() -> None:
    payload = _build_docx_payload_missing_image()
    blob = _inline_blob_from_payload(payload)
    document = _DocumentStub(
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        blob=blob,
    )
    parser = DocxDocumentParser()

    with pytest.raises(ValueError, match="docx_image_missing"):
        parser.parse(document, config={})


def test_docx_parser_uses_alt_text_as_context_fallback() -> None:
    payload, _ = _build_docx_payload_with_alt_fallback()
    blob = _inline_blob_from_payload(payload)
    document = _DocumentStub(
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        blob=blob,
    )
    parser = DocxDocumentParser()

    result = parser.parse(document, config={})

    assert result.assets
    asset = result.assets[0]
    assert asset.context_before == "Image intro text."
    assert asset.context_after == "Fallback description"
    candidates = asset.metadata.get("caption_candidates", [])
    assert candidates and candidates[0] == ("alt_text", "Fallback description")


def test_docx_parser_rejects_oversized_assets() -> None:
    payload = _build_docx_payload_with_large_image(25 * 1024 * 1024 + 1)
    blob = _inline_blob_from_payload(payload)
    document = _DocumentStub(
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        blob=blob,
    )
    parser = DocxDocumentParser()

    with pytest.raises(ValueError, match="docx_asset_too_large"):
        parser.parse(document, config={})


def test_docx_parser_enforces_zip_entry_limits() -> None:
    payload = _build_docx_payload_exceeding_entry_limit(10_001)
    blob = _inline_blob_from_payload(payload)
    document = _DocumentStub(
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        blob=blob,
    )
    parser = DocxDocumentParser()

    with pytest.raises(ValueError, match="docx_archive_limits"):
        parser.parse(document, config={})
